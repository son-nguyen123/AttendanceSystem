from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from app.services.bank_payroll import _build_document


class BankPayrollDocumentTests(unittest.TestCase):
    def test_document_matches_plain_five_column_account_list(self) -> None:
        employees = [
            {
                "employee_code": "1006",
                "name": "Hồ Thị Tú Uyên",
                "account_number": "5601884980",
                "salary": 3330000,
            }
        ]
        with TemporaryDirectory() as temporary_dir:
            output = Path(temporary_dir) / "bank-payroll.docx"
            _build_document(output, employees, "1", "THÁNG 6/2026")
            document = Document(output)

        self.assertEqual(document.paragraphs[0].text, "THIÊN TRÍ THÁNG 6/2026")
        self.assertEqual(len(document.tables), 1)
        table = document.tables[0]
        self.assertEqual(len(table.columns), 5)
        self.assertEqual(
            [cell.text for cell in table.rows[0].cells],
            ["STT", "MÃ NHÂN\nVIÊN", "HỌ TÊN", "SỐ TK", "LƯƠNG"],
        )
        self.assertEqual(
            [cell.text for cell in table.rows[1].cells],
            ["1", "1006", "HO THI TU UYEN", "5601884980", "3,330,000"],
        )
        self.assertEqual(table.rows[-1].cells[0].text, "TỔNG")
        self.assertEqual(table.rows[-1].cells[4].text, "3,330,000")


if __name__ == "__main__":
    unittest.main()
