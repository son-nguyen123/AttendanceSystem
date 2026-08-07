from __future__ import annotations

import json
import re
import shutil
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from app.services.cloud_sync import get_cloud_config
from app.services.bank_account_store import find_duplicate_account_groups, normalize_account_number
from app.services.payroll_store import normalize_employee_name


STORAGE_DIR = Path(__file__).resolve().parents[2] / "storage"
REGISTRY_PATH = STORAGE_DIR / "bank_accounts.json"
SESSION_DIR = STORAGE_DIR / "bank_payroll"


def scan_official_workbook(path: Path, factory: str) -> dict[str, Any]:
    factory = _factory(factory)
    keep_vba = path.suffix.lower() == ".xlsm"
    values_wb = load_workbook(path, data_only=True, keep_vba=keep_vba)
    # A workbook created by the app can contain formulas without cached
    # results (for example immediately after export). Keep a formula view so
    # the employee rows are still discoverable instead of rejecting the file.
    formula_wb = load_workbook(path, data_only=False, keep_vba=keep_vba)
    employees: dict[str, dict[str, Any]] = {}
    month = year = None
    try:
        for sheet_index, values_ws in enumerate(values_wb.worksheets):
            formula_ws = formula_wb.worksheets[sheet_index] if sheet_index < len(formula_wb.worksheets) else values_ws
            if month is None or year is None:
                month, year = _sheet_period(values_ws)
            for header_row in range(1, values_ws.max_row + 1):
                fields = _header_fields(formula_ws, header_row)
                if "code" not in fields or "final_salary" not in fields:
                    continue
                result_row = _find_result_row(values_ws, header_row, fields)
                if result_row is None:
                    result_row = _find_result_row(formula_ws, header_row, fields)
                if result_row is None:
                    continue
                code = _employee_code(
                    values_ws.cell(result_row, fields["code"]).value
                    or formula_ws.cell(result_row, fields["code"]).value
                )
                if not code:
                    continue
                name_col = fields.get("name", fields["code"] + 1)
                name = str(
                    values_ws.cell(result_row, name_col).value
                    or formula_ws.cell(result_row, name_col).value
                    or ""
                )
                name = normalize_employee_name(name)
                if not name:
                    name = normalize_employee_name(_find_employee_name(values_ws, result_row, fields))
                salary = _number(values_ws.cell(result_row, fields["final_salary"]).value)
                calculated_salary = _calculate_salary(values_ws, result_row, fields)
                # Excel may leave formula cells without cached results (or
                # with a stale zero). Rebuild from visible payroll inputs in
                # that case so the bank list does not show false zeroes.
                if salary is None or (salary == 0 and calculated_salary > 0):
                    salary = calculated_salary
                work_days = _work_days(values_ws, result_row, fields)
                header = str(values_ws.cell(header_row, fields["final_salary"]).value or "")
                found = re.search(r"(\d{1,2})\D+(20\d{2})", header)
                if found:
                    month, year = int(found.group(1)), int(found.group(2))
                employees[code] = {
                    "employee_code": code,
                    "name": name,
                    "salary": round(float(salary or 0)),
                    "work_days": round(float(work_days or 0), 2),
                }
    finally:
        values_wb.close()
        formula_wb.close()

    if not employees:
        raise ValueError("Không nhận ra danh sách lương trong file. Hãy chọn đúng bảng chính thức Output 2.")

    # Only people with recorded work in this month's payroll belong in the
    # transfer list.  Keep the old all-rows behaviour as a safe fallback for
    # legacy workbooks that do not expose a usable work-day column.
    active = [item for item in employees.values() if item["work_days"] > 0]
    if active:
        employees = {item["employee_code"]: item for item in active}

    registry = _load_registry_with_drive_fallback(factory)
    duplicate_codes: dict[str, list[str]] = {}
    for group in find_duplicate_account_groups(factory, registry):
        codes = [str(code) for code in group.get("employee_codes") or []]
        for code in codes:
            duplicate_codes[code] = [other for other in codes if other != code]
    for item in employees.values():
        saved = registry.get(_key(factory, item["employee_code"]), {})
        item["account_number"] = normalize_account_number(saved.get("account_number"))
        item["conflict_accounts"] = list(saved.get("conflict_accounts") or [])
        item["conflict_codes"] = duplicate_codes.get(item["employee_code"], [])

    scan_id = datetime.now().strftime("%Y%m%d%H%M%S%f")
    SESSION_DIR.mkdir(parents=True, exist_ok=True)
    snapshot = {
        "scan_id": scan_id,
        "factory": factory,
        "month": month,
        "year": year,
        "source_filename": path.name,
        "employees": list(employees.values()),
    }
    _atomic_json(SESSION_DIR / f"{scan_id}.json", snapshot)
    return snapshot


def list_accounts(factory: str) -> dict[str, Any]:
    factory = _factory(factory)
    registry = _load_registry()
    duplicate_codes: dict[str, list[str]] = {}
    for group in find_duplicate_account_groups(factory, registry):
        codes = [str(code) for code in group.get("employee_codes") or []]
        for code in codes:
            duplicate_codes[code] = [other for other in codes if other != code]
    records = []
    for key, value in registry.items():
        if not isinstance(value, dict) or value.get("factory") != factory:
            continue
        code = _employee_code(value.get("employee_code") or key.rsplit(":", 1)[-1])
        records.append({**value, "conflict_codes": duplicate_codes.get(code, [])})
    return {"factory": factory, "accounts": sorted(records, key=lambda row: _code_sort(row["employee_code"]))}


def save_accounts(factory: str, accounts: list[dict[str, Any]]) -> dict[str, Any]:
    factory = _factory(factory)
    registry = _load_registry()
    updated = 0
    next_registry = dict(registry)
    for row in accounts:
        code = _employee_code(row.get("employee_code"))
        if not code:
            continue
        raw_account = re.sub(r"\s+", "", str(row.get("account_number") or ""))
        account = normalize_account_number(raw_account)
        if raw_account and not account:
            raise ValueError(f"Số tài khoản của mã {code} phải có từ 8 đến 20 chữ số.")
        next_registry[_key(factory, code)] = {
            "factory": factory,
            "employee_code": code,
            "name": str(row.get("name") or "").strip(),
            "account_number": account,
            "conflict_accounts": list(row.get("conflict_accounts") or []),
            "updated_at": datetime.now().isoformat(timespec="seconds"),
        }
        updated += 1
    duplicate_groups = find_duplicate_account_groups(factory, next_registry)
    if duplicate_groups:
        details = "; ".join(
            f"{group['account_number']}: mã {', '.join(group['employee_codes'])}"
            for group in duplicate_groups[:8]
        )
        raise ValueError("Số tài khoản đang trùng giữa các mã, cần kiểm tra lại: " + details)
    registry = next_registry
    _atomic_json(REGISTRY_PATH, registry)
    return {"status": "ok", "updated": updated}


def import_accounts_from_word(
    path: Path,
    factory: str,
    month: int,
    year: int,
    mode: str = "fill_missing",
) -> dict[str, Any]:
    factory = _factory(factory)
    detected_month, detected_year = detect_word_period(path)
    month = detected_month or month
    year = detected_year or year
    if not 1 <= month <= 12 or year < 2000:
        raise ValueError("Tháng hoặc năm lưu không hợp lệ.")
    document = Document(path)
    found: dict[str, dict[str, str]] = {}
    for table in document.tables:
        if not table.rows:
            continue
        header_row = None
        columns: dict[str, int] = {}
        for row_index, row in enumerate(table.rows[:4]):
            labels = [_normalize(cell.text) for cell in row.cells]
            code_col = next((i for i, label in enumerate(labels) if label in {"ma", "ma nhan vien"}), None)
            account_col = next((i for i, label in enumerate(labels) if label in {"so tk", "so tai khoan", "stk"}), None)
            if code_col is not None and account_col is not None:
                name_col = next((i for i, label in enumerate(labels) if "ho ten" in label or "ho va ten" in label), None)
                header_row = row_index
                columns = {"code": code_col, "account": account_col}
                if name_col is not None:
                    columns["name"] = name_col
                break
        if header_row is None:
            continue
        for row in table.rows[header_row + 1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if max(columns.values()) >= len(cells):
                continue
            code = _employee_code(cells[columns["code"]])
            account = normalize_account_number(re.sub(r"\D", "", cells[columns["account"]]))
            if not code or not account:
                continue
            found.setdefault(code, {})[account] = cells[columns.get("name", columns["code"])]

    if not found:
        raise ValueError("Không tìm thấy cột Mã nhân viên và Số tài khoản trong file Word.")

    mode = "replace" if mode == "replace" else "fill_missing"
    registry = _load_registry()
    conflicts = []
    skipped_existing = []
    updated = 0
    source_duplicates = _duplicate_source_accounts(found)
    for code, account_names in found.items():
        key = _key(factory, code)
        current = registry.get(key, {})
        accounts = set(account_names)
        current_account = normalize_account_number(current.get("account_number"))
        if current_account and mode == "fill_missing":
            accounts.add(current_account)
        ordered = sorted(accounts)
        conflict_accounts = ordered if len(ordered) > 1 else []
        if conflict_accounts:
            conflicts.append({
                "employee_code": code,
                "accounts": conflict_accounts,
                "existing_account": current_account,
                "existing_name": str(current.get("name") or "").strip(),
                "file_accounts": sorted(account_names),
                "name": str(next(iter(account_names.values()), "") or "").strip(),
            })
            continue
        if code in source_duplicates:
            conflicts.append({
                "employee_code": code,
                "accounts": ordered,
                "duplicate_codes": source_duplicates[code],
                "reason": "duplicate_account",
                "existing_account": current_account,
                "existing_name": str(current.get("name") or "").strip(),
                "file_accounts": sorted(account_names),
                "name": str(next(iter(account_names.values()), "") or "").strip(),
            })
            continue
        if current_account and mode == "fill_missing":
            skipped_existing.append(code)
            continue
        if not ordered:
            continue
        existing_owners = [owner for owner in _account_codes(registry, factory, ordered[0]) if owner != code]
        if existing_owners:
            conflicts.append({
                "employee_code": code,
                "accounts": ordered,
                "duplicate_codes": existing_owners,
                "reason": "duplicate_account",
                "existing_account": current_account,
                "existing_name": str(current.get("name") or "").strip(),
                "file_accounts": sorted(account_names),
                "name": str(next(iter(account_names.values()), "") or "").strip(),
            })
            continue
        registry[key] = {
            "factory": factory,
            "employee_code": code,
            "name": str((next(iter(account_names.values()), "") if mode == "replace" else current.get("name")) or ""),
            "account_number": ordered[0],
            "conflict_accounts": conflict_accounts,
            "updated_at": datetime.now().isoformat(timespec="seconds"),
            "source_period": f"{year}-{month:02d}",
        }
        updated += 1
    _atomic_json(REGISTRY_PATH, registry)

    config = get_cloud_config()
    drive_path = None
    if config.get("drive_backup_enabled"):
        target_dir = Path(config["drive_root_path"]) / "DuLieuNganHang" / f"{year}-{month:02d}"
        target_dir.mkdir(parents=True, exist_ok=True)
        drive_path = target_dir / f"Xuong{2 if factory == 'factory2' else 1}_{year}-{month:02d}_DanhSachTaiKhoan.docx"
        if path.resolve() != drive_path.resolve():
            shutil.copy2(path, drive_path)
        factory_dir = Path(config["drive_root_path"]) / "DuLieuNganHang" / f"Xuong{2 if factory == 'factory2' else 1}"
        root_registry = factory_dir / "DanhSachTaiKhoan.json"
        root_registry.parent.mkdir(parents=True, exist_ok=True)
        _atomic_json(root_registry, _registry_for_factory(factory))

    return {
        "status": "ok",
        "imported": updated,
        "conflicts": conflicts,
        "skipped_existing": skipped_existing,
        "mode": mode,
        "drive_path": str(drive_path) if drive_path else None,
        "month": month,
        "year": year,
    }


def import_accounts_from_excel_salary(
    path: Path,
    factory: str,
    month: int,
    year: int,
    mode: str = "fill_missing",
) -> dict[str, Any]:
    """Import employee/account pairs from the old salary Excel workbook.

    The old salary template is not the official Output 2 layout.  It has a
    small table headed by ``Mã Nhân Viên`` and ``SỐ TK``.  Some workbooks also
    contain rows from another month, so a row period in ``NỘI DUNG`` is used
    as a filter when it is present.
    """
    factory = _factory(factory)
    detected_month, detected_year = detect_excel_salary_period(path)
    month = detected_month or month
    year = detected_year or year
    if not 1 <= month <= 12 or year < 2000:
        raise ValueError("Tháng hoặc năm lưu không hợp lệ.")

    found: dict[str, dict[str, str]] = {}
    keep_vba = path.suffix.lower() == ".xlsm"
    workbook = load_workbook(path, data_only=True, keep_vba=keep_vba)
    try:
        for worksheet in workbook.worksheets:
            header_row, columns = _excel_salary_columns(worksheet)
            if header_row is None:
                continue
            for row in worksheet.iter_rows(min_row=header_row + 1, values_only=True):
                if not row or max(columns.values()) >= len(row):
                    continue
                code = _employee_code(row[columns["code"]])
                if not code or _normalize(code) in {"stt", "tong", "total"}:
                    continue
                account = normalize_account_number(re.sub(r"\D", "", str(row[columns["account"]] or "")))
                if not account:
                    continue
                row_period = _period_from_text(row[columns["period"]]) if "period" in columns else (None, None)
                if row_period != (None, None) and row_period != (month, year):
                    continue
                name = normalize_employee_name(row[columns.get("name", columns["code"])])
                found.setdefault(code, {})[account] = name
    finally:
        workbook.close()

    if not found:
        raise ValueError("Không tìm thấy cột Mã Nhân Viên và Số TK trong file Excel lương cũ.")

    return _persist_imported_accounts(
        found,
        factory=factory,
        month=month,
        year=year,
        mode=mode,
        source_path=path,
        source_type="excel_salary",
    )


def detect_excel_salary_period(path: Path) -> tuple[int | None, int | None]:
    keep_vba = path.suffix.lower() == ".xlsm"
    workbook = load_workbook(path, data_only=True, keep_vba=keep_vba)
    texts: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        texts.append(str(cell.value))
    finally:
        workbook.close()

    for text in texts:
        month, year = _period_from_text(text)
        if month and year:
            return month, year
    filename_match = re.search(r"(?:thang|t)[_ -]?(\d{1,2})[_ -/]+(20\d{2})", path.stem, re.IGNORECASE)
    if filename_match:
        month, year = int(filename_match.group(1)), int(filename_match.group(2))
        if 1 <= month <= 12:
            return month, year
    return None, None


def _excel_salary_columns(worksheet) -> tuple[int | None, dict[str, int]]:
    for row_index in range(1, min(worksheet.max_row, 30) + 1):
        labels = [_normalize(worksheet.cell(row_index, column).value) for column in range(1, worksheet.max_column + 1)]
        code_column = next((index for index, label in enumerate(labels) if label in {"ma", "ma nhan vien", "ma nv"}), None)
        account_column = next((index for index, label in enumerate(labels) if label in {"so tk", "so tai khoan", "stk"}), None)
        if code_column is None or account_column is None:
            continue
        name_column = next((index for index, label in enumerate(labels) if "ho ten" in label or "ten nhan vien" in label), None)
        period_column = next((index for index, label in enumerate(labels) if "noi dung" in label or label == "thang"), None)
        columns = {"code": code_column, "account": account_column}
        if name_column is not None:
            columns["name"] = name_column
        if period_column is not None:
            columns["period"] = period_column
        return row_index, columns
    return None, {}


def _duplicate_source_accounts(found: dict[str, dict[str, str]]) -> dict[str, list[str]]:
    by_account: dict[str, list[str]] = {}
    for code, account_names in found.items():
        for account in account_names:
            by_account.setdefault(account, []).append(code)
    result: dict[str, list[str]] = {}
    for codes in by_account.values():
        unique_codes = sorted(set(codes), key=_code_sort)
        if len(unique_codes) < 2:
            continue
        for code in unique_codes:
            result[code] = [other for other in unique_codes if other != code]
    return result


def _account_codes(registry: dict[str, dict[str, Any]], factory: str, account: str) -> list[str]:
    owners: list[str] = []
    normalized_factory = _factory(factory)
    for key, value in registry.items():
        if not isinstance(value, dict) or _factory(value.get("factory") or key.split(":", 1)[0]) != normalized_factory:
            continue
        if normalize_account_number(value.get("account_number")) != account:
            continue
        code = _employee_code(value.get("employee_code") or key.rsplit(":", 1)[-1])
        if code and code not in owners:
            owners.append(code)
    return owners


def _persist_imported_accounts(
    found: dict[str, dict[str, str]],
    *,
    factory: str,
    month: int,
    year: int,
    mode: str,
    source_path: Path,
    source_type: str,
) -> dict[str, Any]:
    mode = "replace" if mode == "replace" else "fill_missing"
    registry = _load_registry()
    conflicts = []
    skipped_existing = []
    updated = 0
    source_duplicates = _duplicate_source_accounts(found)
    for code, account_names in found.items():
        key = _key(factory, code)
        current = registry.get(key, {})
        accounts = set(account_names)
        current_account = normalize_account_number(current.get("account_number"))
        if current_account and mode == "fill_missing":
            accounts.add(current_account)
        ordered = sorted(accounts)
        conflict_accounts = ordered if len(ordered) > 1 else []
        if conflict_accounts:
            conflicts.append({
                "employee_code": code,
                "accounts": conflict_accounts,
                "existing_account": current_account,
                "existing_name": str(current.get("name") or "").strip(),
                "file_accounts": sorted(account_names),
                "name": str(next(iter(account_names.values()), "") or "").strip(),
            })
            continue
        if code in source_duplicates:
            conflicts.append({
                "employee_code": code,
                "accounts": ordered,
                "duplicate_codes": source_duplicates[code],
                "reason": "duplicate_account",
                "existing_account": current_account,
                "existing_name": str(current.get("name") or "").strip(),
                "file_accounts": sorted(account_names),
                "name": str(next(iter(account_names.values()), "") or "").strip(),
            })
            continue
        if current_account and mode == "fill_missing":
            skipped_existing.append(code)
            continue
        if not ordered:
            continue
        existing_owners = [owner for owner in _account_codes(registry, factory, ordered[0]) if owner != code]
        if existing_owners:
            conflicts.append({
                "employee_code": code,
                "accounts": ordered,
                "duplicate_codes": existing_owners,
                "reason": "duplicate_account",
                "existing_account": current_account,
                "existing_name": str(current.get("name") or "").strip(),
                "file_accounts": sorted(account_names),
                "name": str(next(iter(account_names.values()), "") or "").strip(),
            })
            continue
        registry[key] = {
            "factory": factory,
            "employee_code": code,
            "name": str((next(iter(account_names.values()), "") if mode == "replace" else current.get("name")) or ""),
            "account_number": ordered[0],
            "conflict_accounts": conflict_accounts,
            "updated_at": datetime.now().isoformat(timespec="seconds"),
            "source_period": f"{year}-{month:02d}",
        }
        updated += 1
    _atomic_json(REGISTRY_PATH, registry)

    config = get_cloud_config()
    drive_path = None
    if config.get("drive_backup_enabled"):
        target_dir = Path(config["drive_root_path"]) / "DuLieuNganHang" / f"{year}-{month:02d}"
        target_dir.mkdir(parents=True, exist_ok=True)
        suffix = source_path.suffix.lower() if source_path.suffix.lower() in {".xlsx", ".xlsm"} else ".xlsx"
        drive_path = target_dir / f"Xuong{2 if factory == 'factory2' else 1}_{year}-{month:02d}_DanhSachTaiKhoan{suffix}"
        if source_path.resolve() != drive_path.resolve():
            shutil.copy2(source_path, drive_path)
        factory_dir = Path(config["drive_root_path"]) / "DuLieuNganHang" / f"Xuong{2 if factory == 'factory2' else 1}"
        root_registry = factory_dir / "DanhSachTaiKhoan.json"
        root_registry.parent.mkdir(parents=True, exist_ok=True)
        _atomic_json(root_registry, _registry_for_factory(factory))

    return {
        "status": "ok",
        "imported": updated,
        "conflicts": conflicts,
        "skipped_existing": skipped_existing,
        "mode": mode,
        "drive_path": str(drive_path) if drive_path else None,
        "month": month,
        "year": year,
        "source_type": source_type,
    }


def _period_from_text(value: Any) -> tuple[int | None, int | None]:
    normalized = _normalize(value)
    patterns = (
        r"(?:thang|t)\s*(\d{1,2})\s*[/.-]\s*(20\d{2})",
        r"(\d{1,2})\s*[/.-]\s*(20\d{2})",
    )
    for pattern in patterns:
        match = re.search(pattern, normalized, re.IGNORECASE)
        if match:
            month, year = int(match.group(1)), int(match.group(2))
            if 1 <= month <= 12:
                return month, year
    return None, None


def detect_word_period(path: Path) -> tuple[int | None, int | None]:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows[:5]:
            texts.extend(cell.text for cell in row.cells)
    combined = " ".join(texts)
    patterns = (
        r"th[aá]ng\s*(\d{1,2})\s*[/.-]\s*(20\d{2})",
        r"(\d{1,2})\s*[/.-]\s*(20\d{2})",
    )
    normalized = combined.lower()
    for pattern in patterns:
        match = re.search(pattern, normalized, re.IGNORECASE)
        if match:
            month, year = int(match.group(1)), int(match.group(2))
            if 1 <= month <= 12:
                return month, year
    filename_match = re.search(r"(?:thang[_ -]?)?(\d{1,2})[_ -]+(20\d{2})", path.stem, re.IGNORECASE)
    if filename_match:
        month, year = int(filename_match.group(1)), int(filename_match.group(2))
        if 1 <= month <= 12:
            return month, year
    return None, None


def export_bank_docx(scan_id: str, account_overrides: list[dict[str, Any]] | None = None) -> tuple[Path, str]:
    snapshot_path = SESSION_DIR / f"{scan_id}.json"
    if not snapshot_path.exists():
        raise ValueError("Phiên quét bảng lương không còn tồn tại. Hãy quét lại file.")
    snapshot = json.loads(snapshot_path.read_text(encoding="utf-8"))
    factory = snapshot["factory"]
    registry = _load_registry()
    overrides = {
        _employee_code(row.get("employee_code")): normalize_account_number(row.get("account_number"))
        for row in (account_overrides or [])
        if _employee_code(row.get("employee_code"))
    }
    employees = []
    missing = []
    seen_accounts: dict[str, str] = {}
    duplicates = []
    for item in snapshot["employees"]:
        saved = registry.get(_key(factory, item["employee_code"]), {})
        account = overrides.get(item["employee_code"], normalize_account_number(saved.get("account_number")))
        if not account:
            missing.append(item["employee_code"])
        elif account in seen_accounts:
            duplicates.append(f"{seen_accounts[account]} và {item['employee_code']}")
        else:
            seen_accounts[account] = item["employee_code"]
        employees.append({**item, "account_number": account})
    if missing:
        raise ValueError("Chưa có số tài khoản cho mã: " + ", ".join(missing[:12]))
    if duplicates:
        raise ValueError("Phát hiện số tài khoản trùng ở mã: " + ", ".join(duplicates[:8]))

    month = snapshot.get("month")
    year = snapshot.get("year")
    factory_no = "2" if factory == "factory2" else "1"
    period = f"THÁNG {month}/{year}" if month and year else "BẢNG LƯƠNG NGÂN HÀNG"
    filename = f"Xuong{factory_no}_{year or 'KhongRo'}-{int(month or 0):02d}_BangLuongNganHang.docx"
    output = SESSION_DIR / f"{scan_id}_{filename}"
    _build_document(output, employees, factory_no, period)
    return output, filename


def export_bank_excel(scan_id: str, account_overrides: list[dict[str, Any]] | None = None) -> tuple[Path, str]:
    """Export the scanned payroll rows as the bank's Excel payment layout.

    The account validation is intentionally the same as the legacy Word
    export.  The resulting workbook follows the user's bank template:
    employee code, name, account number, transfer amount and payment note.
    """
    snapshot_path = SESSION_DIR / f"{scan_id}.json"
    if not snapshot_path.exists():
        raise ValueError("PhiÃªn quÃ©t báº£ng lÆ°Æ¡ng khÃ´ng cÃ²n tá»“n táº¡i. HÃ£y quÃ©t láº¡i file.")
    snapshot = json.loads(snapshot_path.read_text(encoding="utf-8"))
    factory = snapshot["factory"]
    registry = _load_registry()
    overrides = {
        _employee_code(row.get("employee_code")): normalize_account_number(row.get("account_number"))
        for row in (account_overrides or [])
        if _employee_code(row.get("employee_code"))
    }
    employees = []
    missing = []
    seen_accounts: dict[str, str] = {}
    duplicates = []
    for item in snapshot["employees"]:
        saved = registry.get(_key(factory, item["employee_code"]), {})
        account = overrides.get(item["employee_code"], normalize_account_number(saved.get("account_number")))
        if not account:
            missing.append(item["employee_code"])
        elif account in seen_accounts:
            duplicates.append(f"{seen_accounts[account]} vÃ  {item['employee_code']}")
        else:
            seen_accounts[account] = item["employee_code"]
        employees.append({**item, "account_number": account})
    if missing:
        raise ValueError("ChÆ°a cÃ³ sá»‘ tÃ i khoáº£n cho mÃ£: " + ", ".join(missing[:12]))
    if duplicates:
        raise ValueError("PhÃ¡t hiá»‡n sá»‘ tÃ i khoáº£n trÃ¹ng á»Ÿ mÃ£: " + ", ".join(duplicates[:8]))

    month = snapshot.get("month")
    year = snapshot.get("year")
    factory_no = "2" if factory == "factory2" else "1"
    filename = f"Xuong{factory_no}_{year or 'KhongRo'}-{int(month or 0):02d}_BangLuongNganHang.xlsx"
    output = SESSION_DIR / f"{scan_id}_{filename}"
    _build_bank_workbook(output, employees, month, year)
    return output, filename


def export_bank_excel_local(
    scan_id: str,
    account_overrides: list[dict[str, Any]] | None = None,
) -> tuple[Path, str]:
    """Create the bank workbook and copy it to the configured local folder."""

    source_path, filename = export_bank_excel(scan_id, account_overrides)
    config = get_cloud_config()
    raw_directory = str(config.get("local_export_dir") or (STORAGE_DIR / "local_exports"))
    target_directory = Path(raw_directory).expanduser()
    if not target_directory.is_absolute():
        target_directory = (STORAGE_DIR / target_directory).resolve()
    target_directory.mkdir(parents=True, exist_ok=True)
    target_path = target_directory / filename
    shutil.copy2(source_path, target_path)
    return target_path, filename


def backup_registry_to_drive(factory: str = "factory1") -> dict[str, Any]:
    config = get_cloud_config()
    if not config.get("drive_backup_enabled"):
        raise ValueError("Chưa bật thư mục Google Drive trong mục Cloud.")
    if not REGISTRY_PATH.exists():
        _atomic_json(REGISTRY_PATH, {})
    factory = _factory(factory)
    target_dir = Path(config["drive_root_path"]) / "DuLieuNganHang" / f"Xuong{2 if factory == 'factory2' else 1}"
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / "DanhSachTaiKhoan.json"
    _atomic_json(target, _registry_for_factory(factory))
    return {"status": "ok", "path": str(target), "factory": factory}


def restore_registry_from_drive(factory: str = "factory1") -> dict[str, Any]:
    config = get_cloud_config()
    bank_dir = Path(config["drive_root_path"]) / "DuLieuNganHang"
    factory_no = "2" if _factory(factory) == "factory2" else "1"
    word_files = sorted(
        bank_dir.rglob(f"Xuong{factory_no}_*_DanhSachTaiKhoan.docx"),
        key=lambda item: item.stat().st_mtime,
        reverse=True,
    ) if bank_dir.exists() else []
    if word_files:
        source_word = word_files[0]
        month, year = detect_word_period(source_word)
        if month and year:
            result = import_accounts_from_word(source_word, factory, month, year)
            return {
                "status": "ok",
                "restored": result["imported"],
                "conflicts": result["conflicts"],
                "path": str(source_word),
                "month": month,
                "year": year,
                "source": "latest_word",
            }

    source = bank_dir / f"Xuong{2 if _factory(factory) == 'factory2' else 1}" / "DanhSachTaiKhoan.json"
    if not source.exists():
        # Backward compatibility with the former shared Drive backup.
        source = bank_dir / "DanhSachTaiKhoan.json"
    if not source.exists():
        raise ValueError("Không tìm thấy bản sao danh sách tài khoản trên Drive.")
    data = json.loads(source.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Bản sao danh sách tài khoản không hợp lệ.")
    current = _load_registry()
    selected_factory = _factory(factory)
    for key, value in data.items():
        if not isinstance(value, dict):
            continue
        code = _employee_code(value.get("employee_code") or key.rsplit(":", 1)[-1])
        if code:
            current[_key(selected_factory, code)] = {**value, "factory": selected_factory, "employee_code": code}
    _atomic_json(REGISTRY_PATH, current)
    return {"status": "ok", "restored": len(data), "path": str(source), "factory": selected_factory}


def bank_status(factory: str = "factory1") -> dict[str, Any]:
    factory = _factory(factory)
    config = get_cloud_config()
    drive_path = Path(config["drive_root_path"]) / "DuLieuNganHang" / f"Xuong{2 if factory == 'factory2' else 1}" / "DanhSachTaiKhoan.json"
    return {
        "local_count": len(_registry_for_factory(factory)),
        "factory": factory,
        "drive_enabled": bool(config.get("drive_backup_enabled")),
        "drive_path": str(drive_path),
        "drive_copy_exists": drive_path.exists(),
    }


def _registry_for_factory(factory: str) -> dict[str, dict[str, Any]]:
    selected = _factory(factory)
    return {
        key: value
        for key, value in _load_registry().items()
        if isinstance(value, dict) and value.get("factory") == selected
    }


def _build_document(path: Path, employees: list[dict[str, Any]], factory_no: str, period: str) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"THIÊN TRÍ {period.upper()}")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(15)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    # Keep the five-column layout within the printable A4 width.  The name
    # column remains the widest one while the account and salary columns have
    # enough room for their right-aligned numeric values.
    widths = [Cm(0.9), Cm(2.2), Cm(6.7), Cm(4.3), Cm(3.9)]
    headers = ["STT", "MÃ NHÂN\nVIÊN", "HỌ TÊN", "SỐ TK", "LƯƠNG"]
    for column, width in zip(table.columns, widths):
        column.width = width
    for index, (cell, title) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = widths[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _cell_margins(cell, 40, 15, 40, 15)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(9)
    _repeat_header(table.rows[0])

    for index, employee in enumerate(employees, start=1):
        row = table.add_row()
        _prevent_row_split(row)
        cells = row.cells
        values = [
            str(index),
            # Employee codes are identifiers, not amounts: never add a
            # thousands separator (e.g. 1006 must stay 1006).
            _employee_code(employee["employee_code"]),
            normalize_employee_name(employee["name"]),
            employee["account_number"],
            f"{(_number(employee.get('salary')) or 0):,.0f}",
        ]
        for col, (cell, value) in enumerate(zip(cells, values)):
            cell.width = widths[col]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _cell_margins(cell, 40, 10, 40, 10)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if col == 2
                else WD_ALIGN_PARAGRAPH.RIGHT
                if col == 4
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(9)

    total_cells = table.add_row().cells
    merged = total_cells[0].merge(total_cells[3])
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("TỔNG")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(9)
    total_cells[4].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    total_paragraph = total_cells[4].paragraphs[0]
    total_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_paragraph.paragraph_format.space_after = Pt(0)
    total_run = total_paragraph.add_run(
        f"{sum((_number(employee.get('salary')) or 0) for employee in employees):,.0f}"
    )
    total_run.bold = True
    total_run.font.name = "Times New Roman"
    total_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    total_run.font.size = Pt(9)
    document.save(path)


def _build_bank_workbook(
    path: Path,
    employees: list[dict[str, Any]],
    month: int | None,
    year: int | None,
) -> None:
    """Build the bank Excel file matching the supplied six-column template."""
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "BangLuongNganHang"
    worksheet.sheet_view.showGridLines = True

    worksheet.merge_cells("A2:F2")
    worksheet["A2"] = f" THIỆN TRÍ THÁNG {month}/{year}" if month and year else " BẢNG LƯƠNG NGÂN HÀNG"
    worksheet["A2"].font = Font(name="Arial", size=16, bold=True)
    worksheet["A2"].alignment = Alignment(horizontal="center", vertical="center")
    worksheet.row_dimensions[2].height = 34.5

    headers = ["STT", "Mã Nhân Viên", "HỌ TÊN", "SỐ TK", "SỐ TIỀN", "NỘI DUNG"]
    for column, title in enumerate(headers, start=1):
        cell = worksheet.cell(row=3, column=column, value=title)
        cell.font = Font(name="Arial", size=10, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.fill = PatternFill("solid", fgColor="FFFF00")

    column_widths = {"A": 8, "B": 20, "C": 36, "D": 22, "E": 18, "F": 18}
    for column, width in column_widths.items():
        worksheet.column_dimensions[column].width = width
    worksheet.row_dimensions[3].height = 30

    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in worksheet.iter_rows(min_row=3, max_row=3, min_col=1, max_col=6):
        for cell in row:
            cell.border = border

    content = f"T{month}/{year}" if month and year else ""
    total = 0
    for index, employee in enumerate(employees, start=1):
        row_number = index + 3
        salary = _number(employee.get("salary")) or 0
        total += salary
        values = [
            index,
            _employee_code(employee.get("employee_code")),
            normalize_employee_name(employee.get("name")),
            normalize_account_number(employee.get("account_number")),
            salary,
            content,
        ]
        for column, value in enumerate(values, start=1):
            cell = worksheet.cell(row=row_number, column=column, value=value)
            cell.border = border
            cell.font = Font(name="Arial", size=10)
            cell.alignment = Alignment(
                horizontal="left" if column == 3 else "right" if column == 5 else "center",
                vertical="center",
            )
            if column in {2, 4}:
                cell.number_format = "@"
            elif column == 5:
                cell.number_format = '#,##0'

    total_row = len(employees) + 4
    worksheet.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=4)
    total_label = worksheet.cell(row=total_row, column=1, value="TỔNG")
    total_label.font = Font(name="Arial", size=10, bold=True)
    total_label.alignment = Alignment(horizontal="center", vertical="center")
    total_amount = worksheet.cell(row=total_row, column=5, value=total)
    total_amount.font = Font(name="Arial", size=10, bold=True)
    total_amount.alignment = Alignment(horizontal="right", vertical="center")
    total_amount.number_format = '#,##0'
    for column in range(1, 7):
        worksheet.cell(row=total_row, column=column).border = border

    worksheet.freeze_panes = "A4"
    workbook.save(path)


def _display_employee_code(value: Any) -> str:
    code = str(value or "").strip()
    if code.isdigit():
        return f"{int(code):,}"
    return code


def _find_result_row(ws, header_row: int, fields: dict[str, int]) -> int | None:
    """Find the summary row for one employee block.

    Different exported Output 2 files place the summary four to six rows
    below the header. Looking for the row that contains both a code and a
    numeric payroll value keeps the scanner compatible with all of them.
    """
    code_col = fields["code"]
    salary_col = fields["final_salary"]
    work_days_col = fields.get("work_days")

    # The generated Output 2 layout always places the employee summary five
    # rows below its header. The row immediately above it is reserved for the
    # bank account and may itself contain a numeric value that looks exactly
    # like an employee code. Prefer the structural summary row so an account
    # number can never become a payroll employee.
    preferred_row = header_row + 5
    if preferred_row <= ws.max_row:
        preferred_code = _employee_code(ws.cell(preferred_row, code_col).value)
        if preferred_code and preferred_code.isdigit():
            return preferred_row

    code_rows: list[int] = []
    for row in range(header_row + 1, min(ws.max_row, header_row + 9) + 1):
        code = _employee_code(ws.cell(row, code_col).value)
        if not code or not code.isdigit():
            continue
        code_rows.append(row)
        salary = _number(ws.cell(row, salary_col).value)
        work_days = _number(ws.cell(row, work_days_col).value) if work_days_col else None
        if salary is not None or work_days is not None:
            return row
    # Formula-only files may have no cached numeric values. The numeric code
    # is still a reliable anchor for the employee block; later logic falls
    # back to all detected rows when activity totals cannot be evaluated.
    return code_rows[0] if code_rows else None


def _find_employee_name(ws, row: int, fields: dict[str, int]) -> str:
    name_col = fields.get("name")
    if name_col:
        value = str(ws.cell(row, name_col).value or "").strip()
        if value:
            return value
    code_col = fields["code"]
    salary_col = fields["final_salary"]
    for col in range(code_col + 1, salary_col):
        value = ws.cell(row, col).value
        if isinstance(value, str) and value.strip() and _normalize(value) not in {"ten", "ghi chu"}:
            return value.strip()
    return ""


def _work_days(ws, row: int, fields: dict[str, int]) -> float:
    work_days_col = fields.get("work_days")
    if work_days_col:
        value = _number(ws.cell(row, work_days_col).value)
        if value is not None:
            return value

    # Legacy files may omit "Số Ngày Đi Làm". Their day totals are normally
    # the attendance columns immediately before the payroll summary columns.
    hours = _attendance_sum(ws, row, fields["code"] - 2)
    return hours / 8 if hours else 0


def _attendance_sum(ws, row: int, end_col: int) -> float:
    return sum(
        value
        for col in range(1, max(1, end_col) + 1)
        if (value := _number(ws.cell(row, col).value)) is not None
    )


def _header_fields(ws, row: int) -> dict[str, int]:
    fields: dict[str, int] = {}
    for col in range(1, min(ws.max_column, 80) + 1):
        label = _normalize(ws.cell(row, col).value)
        if label in {"ma", "ma nhan vien", "ma nv"} or label.startswith("ma nhan vien "):
            fields.setdefault("code", col)
        elif ("ten nhan vien" in label or "ho ten" in label or "ghi chu" in label) and "name" not in fields:
            fields["name"] = col
        elif label == "muc luong":
            fields["monthly_salary"] = col
        elif label == "luong 1 ngay cong":
            fields["daily_salary"] = col
        elif label == "luong 1 gio cong":
            fields["hourly_salary"] = col
        elif label.startswith("so ngay") and "di lam" in label:
            fields["work_days"] = col
        elif label == "gio lam them":
            fields["overtime"] = col
        elif label == "thuong":
            fields["bonus"] = col
        elif label == "phat nq":
            fields["penalty"] = col
        elif label == "ung luong":
            fields["advance"] = col
        elif label.startswith("luong thang") or (label.startswith("luong") and "thang" in label):
            fields["final_salary"] = col
    return fields


def _sheet_period(ws) -> tuple[int | None, int | None]:
    for row in range(1, min(ws.max_row, 20) + 1):
        for col in range(1, min(ws.max_column, 20) + 1):
            text = str(ws.cell(row, col).value or "")
            match = re.search(r"(20\d{2})[-/.](\d{1,2})[-/.]\d{1,2}", text)
            if match:
                return int(match.group(2)), int(match.group(1))
    return None, None


def _calculate_salary(values_ws, row: int, fields: dict[str, int]) -> float:
    def value(name: str) -> float:
        col = fields.get(name)
        return (_number(values_ws.cell(row, col).value) or 0) if col else 0
    daily = value("daily_salary")
    hourly = value("hourly_salary")
    work_days = value("work_days")
    if work_days <= 0:
        work_days = _work_days(values_ws, row, fields)
    overtime = value("overtime")
    if overtime <= 0:
        overtime = _attendance_sum(values_ws, row + 1, fields["code"] - 3)
    bonus = value("bonus")
    penalty = value("penalty")
    advance = value("advance")
    if not daily and fields.get("monthly_salary"):
        daily = value("monthly_salary") / 26
    if not hourly and fields.get("monthly_salary"):
        hourly = value("monthly_salary") / 208
    return daily * work_days + overtime * hourly * 1.5 + bonus - penalty - advance


def _load_registry() -> dict[str, dict[str, Any]]:
    if not REGISTRY_PATH.exists():
        return {}
    try:
        value = json.loads(REGISTRY_PATH.read_text(encoding="utf-8"))
        return value if isinstance(value, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def _load_registry_with_drive_fallback(factory: str | None = None) -> dict[str, dict[str, Any]]:
    local = _load_registry()
    try:
        config = get_cloud_config()
        if not config.get("drive_backup_enabled"):
            return local
        bank_dir = Path(config["drive_root_path"]) / "DuLieuNganHang"
        if factory:
            selected = _factory(factory)
            factory_dir = bank_dir / f"Xuong{2 if selected == 'factory2' else 1}"
            candidates = [factory_dir / "DanhSachTaiKhoan.json", bank_dir / "DanhSachTaiKhoan.json"]
        else:
            candidates = [bank_dir / "DanhSachTaiKhoan.json"]
        merged = dict(local)
        for drive_file in candidates:
            if not drive_file.exists():
                continue
            cloud = json.loads(drive_file.read_text(encoding="utf-8"))
            if not isinstance(cloud, dict):
                continue
            for key, value in cloud.items():
                if not isinstance(value, dict):
                    continue
                code = _employee_code(value.get("employee_code") or key.rsplit(":", 1)[-1])
                if not code:
                    continue
                if factory:
                    selected = _factory(factory)
                    if value.get("factory") and value.get("factory") != selected:
                        continue
                    merged.setdefault(_key(selected, code), {**value, "factory": selected, "employee_code": code})
                else:
                    merged.setdefault(key, value)
        if merged != local:
            _atomic_json(REGISTRY_PATH, merged)
        return merged
    except (OSError, json.JSONDecodeError, KeyError):
        pass
    return local


def _atomic_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(path)


def _factory(value: str) -> str:
    return "factory2" if value == "factory2" else "factory1"


def _key(factory: str, code: str) -> str:
    return f"{factory}:{code}"


def _employee_code(value: Any) -> str:
    if isinstance(value, bool):
        return ""
    if isinstance(value, (int, float)) and float(value).is_integer():
        return str(int(value))
    return str(value or "").strip().replace(",", "")


def _number(value: Any) -> float | None:
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    try:
        return float(str(value or "").replace(",", "").strip())
    except ValueError:
        return None


def _normalize(value: Any) -> str:
    text = unicodedata.normalize("NFD", str(value or "").strip().lower())
    text = text.replace("đ", "d")
    return " ".join("".join(c for c in text if unicodedata.category(c) != "Mn").split())


def _code_sort(value: str) -> tuple[int, str]:
    return (int(value), value) if value.isdigit() else (10**9, value)


def _shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
